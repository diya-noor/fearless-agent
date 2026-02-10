# Fearless Document Formatter - FINAL COMPLETE VERSION
from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import requests
import logging

app = Flask(__name__)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

HEADER_LOGO_URL = "https://raw.githubusercontent.com/diya-noor/fearless-agent/main/fearless_icon_logo.png"
FOOTER_LOGO_URL = "https://raw.githubusercontent.com/diya-noor/fearless-agent/main/fearless_text_logo.png"

COLORS = {
    'orange': RGBColor(238, 83, 64),      # #EE5340 - Title, Heading 2
    'purple': RGBColor(92, 57, 119),      # #5C3977 - Heading 1, Heading 3
    'gray100': RGBColor(73, 79, 86),      # #494F56 - Dark gray (body text)
    'gray70': RGBColor(127, 131, 136),    # #7F8388 - Light gray (subtitle)
}

def download_image(url):
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            return BytesIO(response.content)
    except:
        pass
    return None

def add_header_footer(doc):
    section = doc.sections[0]
    
    # HEADER
    header = section.header
    for para in header.paragraphs:
        para.clear()
    
    header_para = header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_after = Pt(0)
    
    logo_stream = download_image(HEADER_LOGO_URL)
    if logo_stream:
        header_para.add_run().add_picture(logo_stream, height=Inches(0.5))
    
    # FOOTER
    footer = section.footer
    for para in footer.paragraphs:
        para.clear()
    
    logo_para = footer.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_after = Pt(1)
    
    logo_stream = download_image(FOOTER_LOGO_URL)
    if logo_stream:
        logo_para.add_run().add_picture(logo_stream, height=Inches(0.2))
    
    addr_para = footer.add_paragraph()
    addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr_para.paragraph_format.space_before = Pt(0)
    addr_para.paragraph_format.space_after = Pt(0)
    
    run1 = addr_para.add_run("8 Market Place, Suite 200, Baltimore, MD 21202")
    run1.font.name = 'Montserrat'
    run1.font.size = Pt(7)
    run1.font.color.rgb = COLORS['gray100']
    
    contact_para = footer.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    
    run2 = contact_para.add_run("(410) 394-9600  /  fax (410) 779-3706  /  ")
    run2.font.name = 'Montserrat'
    run2.font.size = Pt(7)
    run2.font.color.rgb = COLORS['gray100']
    
    run3 = contact_para.add_run("fearless.tech")
    run3.font.name = 'Montserrat'
    run3.font.size = Pt(7)
    run3.font.color.rgb = COLORS['gray100']

def format_content(doc, text):
    text = text.strip()
    if '\n' not in text and '\\n' in text:
        text = text.replace('\\r\\n', '\n').replace('\\n', '\n')
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        process_paragraph(doc, line)

def process_paragraph(doc, para_text):
    para = doc.add_paragraph()
    
    # More lenient heading detection - allow #Title or # Title
    if para_text.startswith('#'):
        level = 0
        idx = 0
        while idx < len(para_text) and para_text[idx] == '#':
            level += 1
            idx += 1
        
        heading_text = para_text[level:].strip()
        
        if not heading_text:
            # Empty heading - treat as body
            run = para.add_run(para_text)
            run.font.name = 'Montserrat'
            run.font.size = Pt(10)
            run.font.color.rgb = COLORS['gray100']
            para.paragraph_format.space_after = Pt(16)
            return
        
        run = para.add_run(heading_text)
        
        if level == 1:
            # # = TITLE: Orange, LEFT, 28pt, VERY MINIMAL space after
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(32)
            para.paragraph_format.space_after = Pt(2)  # VERY MINIMAL space to subtitle
            run.font.name = 'Montserrat Alternates'
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = COLORS['orange']  # #EE5340 Orange
            logger.info(f"✅ TITLE (Orange, Left): {heading_text}")
            
        elif level == 2:
            # ## = SUBTITLE: Light Gray, Left, 22pt
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(26)
            para.paragraph_format.space_after = Pt(20)
            run.font.name = 'Montserrat Alternates'
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = COLORS['gray70']  # #7F8388 Light Gray
            logger.info(f"✅ SUBTITLE (Light Gray, Left): {heading_text}")
            
        elif level == 3:
            # ### = HEADING 1: PURPLE, Left, 18pt
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(20)
            para.paragraph_format.space_after = Pt(16)
            run.font.name = 'Montserrat'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = COLORS['purple']  # #5C3977 Purple
            logger.info(f"✅ HEADING 1 (Purple, Left): {heading_text}")
            
        elif level == 4:
            # #### = HEADING 2: ORANGE, Left, 14pt
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(12)
            run.font.name = 'Montserrat'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = COLORS['orange']  # #EE5340 Orange
            logger.info(f"✅ HEADING 2 (Orange, Left): {heading_text}")
            
        else:
            # ##### = HEADING 3: PURPLE, Left, 12pt
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(10)
            run.font.name = 'Montserrat'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = COLORS['purple']  # #5C3977 Purple
            logger.info(f"✅ HEADING 3 (Purple, Left): {heading_text}")
    else:
        # Body text: Dark Gray, 10pt, Line height 1.5
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(16)
        para.paragraph_format.line_spacing = 1.5
        
        run = para.add_run(para_text)
        run.font.name = 'Montserrat'
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS['gray100']  # #494F56 Dark Gray
        logger.info(f"📝 Body text: {para_text[:50]}...")

@app.route('/generate-document', methods=['POST'])
def generate_document():
    try:
        logger.info("📝 Generating document...")
        data = request.json
        text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        logger.info(f"Input preview: {text[:200]}...")
        
        doc = Document()
        
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.footer_distance = Inches(0.3)
        
        add_header_footer(doc)
        format_content(doc, text)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        logger.info("✅ Document generated successfully")
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name='fearless_document.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"❌ Error: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'healthy'}), 200

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)